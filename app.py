from flask import Flask, render_template, request, send_file
from werkzeug.utils import secure_filename
import os
from fpdf import FPDF
from docx import Document
from pdf2docx import Converter
from PIL import Image
from io import BytesIO

app = Flask(__name__)

# Set the upload folder and allowed file extensions
app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)  # Ensure uploads folder exists

# Route for the homepage
@app.route('/')
def index():
    return render_template('index.html')

# Route to handle file upload and conversion
@app.route('/convert', methods=['POST'])
def convert():
    file = request.files['file']
    conversion_type = request.form['conversion_type']

    if file:
        # Ensure the upload directory exists
        upload_folder = app.config['UPLOAD_FOLDER']
        os.makedirs(upload_folder, exist_ok=True)  # Create the directory if it doesn't exist
        
        # Save the file
        filename = secure_filename(file.filename)
        file_path = os.path.join(upload_folder, filename)
        file.save(file_path)
        print(f"Saving file to: {file_path}")  # For debugging purposes

        # Handle conversion type
        if conversion_type == 'word-to-pdf':
            output_path = file_path.replace('.docx', '.pdf')
            convert_word_to_pdf(file_path, output_path)
        elif conversion_type == 'pdf-to-word':
            output_path = file_path.replace('.pdf', '.docx')
            convert_pdf_to_word(file_path, output_path)
        else:
            return "Unsupported conversion type", 400

        return send_file(output_path, as_attachment=True)

# Function to convert Word to PDF
def convert_word_to_pdf(input_path, output_path):
    try:
        print(f"Converting Word to PDF: {input_path} -> {output_path}")
        
        # Create FPDF object and set up PDF properties
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Open the .docx file using python-docx
        doc = Document(input_path)
        
        # Set up font for the PDF
        pdf.set_font("Arial", size=12)

        # Loop through each paragraph in the Word document
        for para in doc.paragraphs:
            pdf.add_page()
            
            # Handle special characters by encoding them safely
            text = para.text.encode('latin-1', 'ignore').decode('latin-1')

            # Add the paragraph to the PDF
            pdf.multi_cell(0, 10, text)

        # Handle tables in Word document
        for table in doc.tables:
            pdf.add_page()  # Add a new page for tables
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.encode('latin-1', 'ignore').decode('latin-1')
                    pdf.multi_cell(0, 10, text)
        
        # Handle images in Word document
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image = Image.open(BytesIO(rel.target_part.blob))
                image_path = os.path.join(app.config['UPLOAD_FOLDER'], "image.jpg")
                image.save(image_path)

                # Insert image into PDF
                pdf.add_page()
                pdf.image(image_path, x=10, y=10, w=180)

        # Output PDF
        pdf.output(output_path)

        # Check if the output file exists
        if os.path.exists(output_path):
            print(f"PDF file saved successfully at {output_path}")
        else:
            print(f"Error: PDF file not saved at {output_path}")
    except Exception as e:
        print(f"Error converting Word to PDF: {e}")

# Function to convert PDF to Word
def convert_pdf_to_word(input_path, output_path):
    try:
        converter = Converter(input_path)
        converter.convert(output_path, start=0, end=None)  # Convert the entire PDF
        converter.close()
        print(f"Converted {input_path} to {output_path} successfully.")
    except Exception as e:
        print(f"Error converting PDF to Word: {e}")

if __name__ == '__main__':
    app.run(debug=True)
